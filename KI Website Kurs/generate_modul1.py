from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Seitenränder ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── Farben & Schriften ────────────────────────────────────────
ORANGE  = RGBColor(0xD8, 0x62, 0x2A)
DUNKEL  = RGBColor(0x16, 0x0C, 0x05)
GRAU    = RGBColor(0x7A, 0x6A, 0x5A)
WEISS   = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, size=11, bold=False, color=None, font_name="DM Sans"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color

def heading1(doc, text):
    """Kurs-Titel / Cover-Heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=22, bold=True, color=DUNKEL, font_name="Montserrat")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(doc, text):
    """Modul-Überschrift"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=ORANGE, font_name="Montserrat")
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(doc, text):
    """Video-Überschrift"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=DUNKEL, font_name="Montserrat")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    return p

def meta_line(doc, text):
    """Metainfo (Dauer, Typ)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=9, bold=False, color=GRAU, font_name="Montserrat")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    return p

def body(doc, text):
    """Fliesstext"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, color=DUNKEL)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(17)
    return p

def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 72)
    set_font(run, size=8, color=RGBColor(0xD0, 0xB8, 0x96))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ══════════════════════════════════════════════════════════════
# TITELSEITE
# ══════════════════════════════════════════════════════════════

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("KI WEBSITE WORKSHOP")
set_font(run, size=10, bold=True, color=ORANGE, font_name="Montserrat")
p.paragraph_format.space_after = Pt(8)

heading1(doc, "Modul 1 – Warum das alles verändert")

p = doc.add_paragraph()
run = p.add_run("Ausformulierte Sprechtexte · Stand März 2026")
set_font(run, size=10, color=GRAU, font_name="Montserrat")
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run("3 Videos · ~13 Minuten · Kameraaufnahme")
set_font(run, size=10, color=GRAU, font_name="Montserrat")
p.paragraph_format.space_after = Pt(0)

divider(doc)

p = doc.add_paragraph()
run = p.add_run(
    "Diese Texte sind ausformulierte Sprechversionen der Stichpunkte aus den Aufnahme-Notizen. "
    "Sie dienen als Orientierung und können frei adaptiert werden. "
    "Wörter in [eckigen Klammern] sind Regieanweisungen."
)
set_font(run, size=10, color=GRAU)
p.paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════
# VIDEO 1.1
# ══════════════════════════════════════════════════════════════

doc.add_page_break()

heading2(doc, "VIDEO 1.1")
heading3(doc, "Willkommen – Was dich erwartet")
meta_line(doc, "Kameraaufnahme  ·  ca. 3 Minuten  ·  Keine Bildschirmaufnahme")

divider(doc)

body(doc,
    "Hey. Ich bin Hanna. Und wenn du das hier siehst, hast du eine Entscheidung getroffen, "
    "die ich mir selbst früher sehr gewünscht hätte."
)

body(doc,
    "In diesem Workshop lernst du, wie du mit KI eine professionelle Website baust. "
    "Ohne Code. Ohne Agentur. Ohne monatelangen Abstimmungsprozess."
)

body(doc,
    "Das Ergebnis, das du am Ende in der Hand hältst, ist eine fertige Website. "
    "Eine, die live im Internet ist. Eine, die du selbst gebaut hast. "
    "Und eine, die du ab jetzt eigenständig weiterentwickeln kannst. "
    "Kein vages \"du verstehst KI jetzt ein bisschen besser\" "
    "sondern ein konkretes, sichtbares Ergebnis."
)

body(doc,
    "Der Kurs ist in sechs Module aufgeteilt. Nach jedem einzelnen Modul hast du etwas Handfestes. "
    "Wir reden nicht über Theorie. Wir bauen."
)

body(doc,
    "Ich sage dir das direkt: Dieses Werkzeug hätte ich vor zwei Jahren gebraucht. "
    "Ich habe viel Zeit verschwendet. Agentur, Freelancer, endlose Abstimmungsrunden, Geld. "
    "Mit dem, was du hier lernst, fällt das komplett weg."
)

body(doc,
    "Und das Beste daran: Du brauchst null Vorkenntnisse. "
    "Wenn du eine App auf deinem Handy öffnen kannst, kannst du das hier machen. "
    "Wirklich so einfach."
)

p = doc.add_paragraph()
run = p.add_run("Dann fangen wir an.")
set_font(run, size=12, bold=True, color=DUNKEL)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(0)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("[Energie nach oben, direkter Blick in die Kamera, Video endet hier]")
set_font(run, size=10, color=GRAU, font_name="Montserrat")
p.paragraph_format.space_after = Pt(0)


# ══════════════════════════════════════════════════════════════
# VIDEO 1.2
# ══════════════════════════════════════════════════════════════

doc.add_page_break()

heading2(doc, "VIDEO 1.2")
heading3(doc, "Das Werkzeug, das alles ändert")
meta_line(doc, "Kameraaufnahme + kurze Live-Demo  ·  ca. 5 Minuten")

divider(doc)

body(doc,
    "Bevor wir anfangen zu bauen, müssen wir kurz über das Werkzeug reden. "
    "Nicht weil es Pflicht ist. Sondern weil du, wenn du den Unterschied einmal verstanden hast, "
    "alles was in den nächsten Modulen passiert, sofort einordnen kannst."
)

body(doc,
    "Du kennst ChatGPT. Die meisten von uns haben damit angefangen. "
    "ChatGPT ist gut. ChatGPT erklärt, beantwortet, hilft beim Brainstorming. "
    "Aber ChatGPT hat eine grundlegende Grenze, die die meisten gar nicht bewusst wahrnehmen."
)

body(doc,
    "ChatGPT lebt in einem Chatfenster. Es denkt mit dir, aber es macht nichts. "
    "Du bekommst Text. Was du dann damit anfängst, bist du."
)

body(doc,
    "Claude Code ist etwas grundlegend anderes."
)

body(doc,
    "Claude Code sitzt nicht im Chatfenster. "
    "Claude Code sitzt auf deinem Computer. Und es arbeitet. "
    "Es schreibt Dateien. Es bearbeitet Dateien. Es speichert Dateien. Selbstständig."
)

body(doc,
    "Ich erkläre das gerne mit einem Bild. "
    "ChatGPT ist wie ein Taschenrechner. "
    "Wunderbares Werkzeug, spart Zeit, hilft dir rechnen. "
    "Aber du tippst die Zahlen ein, du schreibst das Ergebnis auf, du machst weiter."
)

body(doc,
    "Claude Code ist der Buchhalter, der die gesamte Steuererklärung für dich macht. "
    "Du sagst: Ich brauche eine fertige Website. Und er baut sie. Komplett. "
    "Du schaust drüber und sagst, was du anders willst. Das war es."
)

body(doc,
    "Das ist kein kleiner Unterschied. Das ist nicht eine Version besser. "
    "Das ist ein fundamental anderes Werkzeug."
)

body(doc,
    "[Wechsel zur kurzen Live-Demo, Bildschirm einblenden oder direkt filmen]"
)

body(doc,
    "Ich zeige dir das jetzt live. Ich öffne Claude Code, tippe einen ganz einfachen Befehl ein. "
    "Erstell mir eine einfache HTML-Seite mit dem Text Hallo Welt."
)

body(doc,
    "[Claude Code tippen, kurze Pause, Datei erscheint im Ordner]"
)

body(doc,
    "Siehst du das? Die Datei ist da. Auf meinem Computer. "
    "Ich habe keine einzige Zeile Code geschrieben. "
    "Ich habe nur gesagt, was ich will."
)

body(doc,
    "[Zurück zur Kamera]"
)

p = doc.add_paragraph()
run = p.add_run(
    "Das ist kein Hype. Das ist kein Marketing. "
    "Das ist wie du ab heute arbeitest."
)
set_font(run, size=12, bold=True, color=DUNKEL)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)

body(doc,
    "Im nächsten Video stelle ich dir noch jemanden vor, der dich durch diesen Kurs begleitet. "
    "Und dann reden wir über das, was dieses Werkzeug wirklich bedeutet."
)


# ══════════════════════════════════════════════════════════════
# VIDEO 1.3
# ══════════════════════════════════════════════════════════════

doc.add_page_break()

heading2(doc, "VIDEO 1.3")
heading3(doc, "Dein Begleiter durch den Kurs – und was das wirklich bedeutet")
meta_line(doc, "Kameraaufnahme · Hanna + Martin gemeinsam oder kurze Vorstellung  ·  ca. 5 Minuten")

divider(doc)

body(doc,
    "Bevor wir in die Installation gehen, möchte ich dir jemanden vorstellen. "
    "Jemanden, der mit mir zusammen durch diesen Kurs führt und der genau das verkörpert, "
    "was Claude Code in der Praxis bedeutet."
)

body(doc,
    "[Martin kurz ins Bild / Martin stellt sich selbst vor]"
)

body(doc,
    "Das ist Martin Thun. Martin ist KI-Unternehmer, Gründer von scale², "
    "und einer der wenigen Menschen, die ich kenne, "
    "der dieses Werkzeug nicht nur versteht, sondern täglich damit baut. "
    "Martin wird dich ab jetzt durch jeden technischen Schritt führen: "
    "von der Installation bis zur fertigen, live geschalteten Website. "
    "Schritt für Schritt, nichts ausgelassen."
)

body(doc,
    "[Zurück zu Hanna oder beide im Bild]"
)

body(doc,
    "Und jetzt möchte ich dir etwas sagen, bevor wir starten. "
    "Etwas, das mir wichtig ist. "
    "Weil ich nicht möchte, dass du diesen Kurs mit falschen Erwartungen beginnst."
)

body(doc,
    "Claude Code ist mächtig. Wirklich mächtig. "
    "Nicht auf die Art wie ein neues Tool mächtig ist, das du zwei Wochen später wieder vergessen hast. "
    "Sondern auf eine Art, die ich als fundamental bezeichnen würde."
)

body(doc,
    "Was wir hier gerade erleben, ist keine Verbesserung von etwas Bestehendem. "
    "Es ist eine Verschiebung. "
    "Software wird seit Jahrzehnten von Entwicklern gebaut. "
    "Menschen, die jahrelang lernen mussten, um eine Datei zu schreiben, eine Seite zu bauen, ein System aufzusetzen. "
    "Claude Code verändert das gerade, während wir zuschauen. "
    "Die gesamte Softwarebranche wird sich in den nächsten Jahren neu erfinden müssen. "
    "Was du hier lernst, ist kein Nischenthema. "
    "Das ist eine der fundamentalsten Verschiebungen, die wir technologisch seit langem gesehen haben."
)

body(doc,
    "Und genau deshalb sage ich dir das nächste jetzt direkt."
)

body(doc,
    "Mit Macht kommt Verantwortung. Das klingt groß, ist aber ganz konkret gemeint. "
    "Dieses Werkzeug kann in kurzer Zeit viel entstehen lassen. "
    "Bitte nutze das bewusst. Nicht um schnell irgendwas rauszuhauen, "
    "sondern um wirklich das zu bauen, was du meinst und was du vertreten kannst. "
    "Eigenverantwortung ist hier kein Lippenbekenntnis, sondern Voraussetzung."
)

body(doc,
    "Und noch eine Sache, die ich dir nicht verschweigen will."
)

body(doc,
    "Die schönste Website, die wir in diesem Kurs bauen, "
    "kann nur so viel leisten wie dein Business trägt. "
    "Das ist keine Einschränkung dieses Kurses. Das ist eine ehrliche Ansage. "
    "Wenn du noch keine Klarheit hast, wen du ansprechen willst, was du anbietest "
    "und warum jemand bei dir kaufen sollte statt woanders, "
    "dann wird eine neue Website das nicht reparieren. "
    "Eine gute Website verstärkt was da ist. Sie ersetzt nicht, was fehlt."
)

body(doc,
    "Ich sage das, weil ich es selbst erlebt habe. "
    "Ich habe Tools genutzt, Systeme gebaut, Seiten gelauncht. "
    "Und der Unterschied zwischen den Momenten wo es geklappt hat und den Momenten wo es nicht geklappt hat "
    "hatte fast nie mit dem Werkzeug zu tun. "
    "Es hatte immer mit der Klarheit dahinter zu tun."
)

body(doc,
    "Also: Wenn du dein Business im Griff hast, wenn du weißt was du tust und für wen, "
    "dann wird das hier ein echter Sprung für dich sein. "
    "Versprochen."
)

p = doc.add_paragraph()
run = p.add_run(
    "Und wenn du gerade noch in diesem Klärungsprozess bist: auch gut. "
    "Dann bau jetzt deine erste Seite, lerne das Werkzeug kennen, "
    "und bring das Business-Fundament parallel dazu in Form."
)
set_font(run, size=11, bold=False, color=DUNKEL)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.line_spacing = Pt(17)

body(doc,
    "Jetzt aber genug geredet. "
    "Martin, du übernimmst."
)

p = doc.add_paragraph()
run = p.add_run("[Übergabe an Martin, Modul 2 beginnt]")
set_font(run, size=10, color=GRAU, font_name="Montserrat")
p.paragraph_format.space_after = Pt(0)


# ══════════════════════════════════════════════════════════════
# SPEICHERN
# ══════════════════════════════════════════════════════════════

output_path = r"G:\Meine Ablage\Claudes Playground\scale²\Client Projects\Hanna KI\Website Workshop\Modul1_Sprechtexte.docx"
doc.save(output_path)
print(f"Gespeichert: {output_path}")
